"""
Генератор Word отчетов
"""

from docx import Document
from docx.shared import Pt
from io import BytesIO
from typing import Dict
from .blocks import CoverPageBlock, SummaryBlock, PlatformBlock


class WordReportGenerator:
    """Генерирует детальные отчеты в формате Word"""

    def __init__(self):
        self.doc = None

    def generate(self, data: Dict) -> BytesIO:
        """
        Генерирует полный отчет в Word

        Args:
            data: Данные аналитики из API

        Returns:
            BytesIO с документом Word
        """
        # Создаем документ
        self.doc = Document()

        # Настройка стилей документа
        self._setup_document_styles()

        # Добавляем титульную страницу
        CoverPageBlock.add_to_document(self.doc, data)

        # Добавляем исполнительное резюме
        SummaryBlock.add_to_document(self.doc, data)

        # Добавляем блоки по каждой платформе
        platforms = data.get("platforms", {})
        for platform_key, platform_data in platforms.items():
            PlatformBlock.add_to_document(self.doc, platform_key, platform_data)

        # Заключение
        self._add_conclusion(data)

        # Сохраняем в BytesIO
        output = BytesIO()
        self.doc.save(output)
        output.seek(0)

        return output

    def _setup_document_styles(self):
        """Настройка стилей документа"""
        # Настройка шрифта для Normal стиля
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(11)

        # Настройка отступов
        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(6)
        paragraph_format.line_spacing = 1.15

    def _add_conclusion(self, data: Dict):
        """Добавляет заключение к отчету"""
        from .utils import format_number, interpret_ps

        self.doc.add_heading("ЗАКЛЮЧЕНИЕ", level=1)

        conclusion_para = self.doc.add_paragraph()

        platforms = data.get("platforms", {})
        total_authors = sum(
            p["aggregated"]["total_authors"] for p in platforms.values()
        )
        total_followers = sum(
            p["aggregated"]["total_followers"] for p in platforms.values()
        )

        # Находим лучшего автора по PS
        best_author = None
        best_ps = 0
        best_platform = None

        for platform_key, platform_data in platforms.items():
            for author in platform_data["authors"]:
                if author["scores"]["PS"] > best_ps:
                    best_ps = author["scores"]["PS"]
                    best_author = author["author_name"]
                    best_platform = platform_key

        conclusion_text = (
            f"Проведенный анализ охватил деятельность {total_authors} "
            f"{'автора' if total_authors == 1 else 'авторов'} "
            f"с общей аудиторией {format_number(total_followers)} подписчиков "
            f"на {len(platforms)} "
            f"{'платформе' if len(platforms) == 1 else 'платформах'}. "
        )

        if best_author:
            from .utils import get_platform_name

            conclusion_text += (
                f"Наиболее высокие показатели эффективности продемонстрировал "
                f"{best_author} на платформе {get_platform_name(best_platform)} "
                f"с Presence Score {best_ps} баллов, что соответствует "
                f'категории "{interpret_ps(best_ps)}".'
            )

        conclusion_text += (
            "\n\nРекомендуется продолжить мониторинг указанных показателей "
            "для выявления долгосрочных трендов и своевременной корректировки "
            "контент-стратегии. Особое внимание следует уделить динамике "
            "Momentum Score как индикатора перспектив роста."
        )

        conclusion_run = conclusion_para.add_run(conclusion_text)
        conclusion_run.font.size = Pt(11)

        self.doc.add_paragraph()

        # Примечания
        notes_heading = self.doc.add_heading("Примечания", level=2)

        notes_para = self.doc.add_paragraph()
        notes_text = (
            "• Presence Score (PS) - интегральный показатель присутствия, "
            "учитывающий охват, вовлеченность и активность публикаций "
            "относительно других авторов на платформе.\n\n"
            "• Momentum Score (MS) - показатель динамики развития, "
            "отражающий темпы роста ключевых метрик в сравнении с конкурентами.\n\n"
            "• Engagement Rate (ER) - уровень вовлеченности аудитории, "
            "рассчитываемый как отношение вовлечений к просмотрам или подписчикам.\n\n"
            "• Share Rate (SR) и Comment Rate (CR) - показатели виральности контента."
        )
        notes_run = notes_para.add_run(notes_text)
        notes_run.font.size = Pt(10)
        notes_run.font.italic = True
