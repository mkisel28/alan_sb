"""
Блок общей сводки по всем платформам
"""

from docx import Document
from docx.shared import Pt, RGBColor
from typing import Dict
from ..utils import (
    format_number,
    format_percent,
    format_delta,
    get_platform_name,
    get_trend_description,
)


class SummaryBlock:
    """Генерирует общую сводку по всем платформам"""

    @staticmethod
    def add_to_document(doc: Document, data: Dict) -> None:
        """
        Добавляет общую сводку в документ

        Args:
            doc: Документ Word
            data: Данные отчета
        """
        # Заголовок сводки
        heading = doc.add_heading("ИСПОЛНИТЕЛЬНОЕ РЕЗЮМЕ", level=1)
        heading_run = heading.runs[0]
        heading_run.font.color.rgb = RGBColor(0, 102, 204)

        doc.add_paragraph()

        # Вводный текст - расширенный
        intro1 = doc.add_paragraph()
        intro1_run = intro1.add_run(
            "Данный отчет представляет собой комплексный и всесторонний анализ эффективности "
            "присутствия в социальных сетях за отчетный период. В рамках настоящего исследования "
            "проводится детальная оценка ключевых показателей производительности (KPI), "
            "анализируется динамика развития аудитории, оценивается уровень вовлеченности "
            "пользователей, а также изучаются тенденции и закономерности, характеризующие "
            "активность на различных цифровых платформах. Отчет подготовлен с учетом лучших "
            "практик анализа социальных медиа и ориентирован на предоставление исчерпывающей "
            "информации для принятия обоснованных стратегических и тактических решений."
        )
        intro1_run.font.size = Pt(11)

        doc.add_paragraph()

        intro2 = doc.add_paragraph()
        intro2_run = intro2.add_run(
            "В процессе подготовки отчета применялись современные методы аналитики социальных "
            "медиа, включая перцентильное ранжирование показателей эффективности, расчет "
            "интегральных метрик (Presence Score и Momentum Score), сравнительный анализ "
            "текущего и предыдущего периодов, а также комплексную оценку динамики изменений. "
            "Использование перцентильного подхода позволяет объективно сравнивать результаты "
            "различных авторов и платформ, учитывая их специфические особенности и масштаб "
            "деятельности. Такой методологический подход обеспечивает высокую точность и "
            "релевантность получаемых выводов."
        )
        intro2_run.font.size = Pt(11)

        doc.add_paragraph()

        intro3 = doc.add_paragraph()
        intro3_run = intro3.add_run(
            "Особое внимание в отчете уделяется сравнительному анализу показателей текущего "
            "и предыдущего периодов. Все данные предыдущего периода в настоящем документе "
            "указываются в отдельной колонке таблиц или в круглых скобках после показателей "
            "текущего периода для удобства восприятия и сравнения. Такой формат представления "
            "информации позволяет наглядно отслеживать тренды, выявлять положительную или "
            "отрицательную динамику по каждому ключевому показателю, а также своевременно "
            "реагировать на изменения в эффективности коммуникационных стратегий. Процентные "
            "изменения и абсолютные дельты рассчитываются автоматически и сопровождаются "
            "качественной интерпретацией для облегчения понимания представленной информации."
        )
        intro3_run.font.size = Pt(11)

        doc.add_paragraph()

        # Агрегированная статистика
        platforms = data.get("platforms", {})

        total_authors = sum(
            p["aggregated"]["total_authors"] for p in platforms.values()
        )
        total_followers = sum(
            p["aggregated"]["total_followers"] for p in platforms.values()
        )
        total_posts = sum(p["aggregated"]["total_posts"] for p in platforms.values())
        total_views = sum(p["aggregated"]["total_views"] for p in platforms.values())
        total_engagement = sum(
            p["aggregated"]["total_engagement"] for p in platforms.values()
        )

        # Средние значения
        avg_ps = (
            sum(p["aggregated"]["avg_PS"] for p in platforms.values()) / len(platforms)
            if platforms
            else 0
        )
        avg_er = (
            sum(p["aggregated"]["avg_ER_view"] for p in platforms.values())
            / len(platforms)
            if platforms
            else 0
        )

        # Ключевые показатели
        doc.add_heading("Ключевые показатели текущего периода", level=2)

        summary_para1 = doc.add_paragraph()
        summary_text1 = (
            f"В течение анализируемого отчетного периода была проведена систематическая "
            f"оценка активности {total_authors} "
            f"{'автора' if total_authors == 1 else 'авторов'} на {len(platforms)} "
            f"{'платформе' if len(platforms) == 1 else 'платформах'} социальных медиа. "
            f"Совокупная аудитория всех исследуемых авторов на момент завершения отчетного "
            f"периода составила {format_number(total_followers)} подписчиков, что представляет "
            f"собой значительный охват целевой аудитории и свидетельствует о существенном "
            f"медийном присутствии в цифровом пространстве. Общее количество контента, "
            f"опубликованного за рассматриваемый период, достигло {format_number(total_posts)} "
            f"{'публикации' if total_posts == 1 else 'публикаций'}, что демонстрирует высокую "
            f"активность авторов и регулярность обновления контента."
        )
        summary_run1 = summary_para1.add_run(summary_text1)
        summary_run1.font.size = Pt(11)

        doc.add_paragraph()

        summary_para2 = doc.add_paragraph()
        summary_text2 = (
            f"Опубликованный контент получил широкий отклик аудитории: суммарное количество "
            f"просмотров по всем платформам и авторам составило {format_number(total_views)}, "
            f"при этом общее число вовлечений (включая лайки, комментарии, репосты и сохранения) "
            f"достигло {format_number(total_engagement)}. Эти показатели характеризуют не только "
            f"количественный охват публикаций, но и качественное взаимодействие аудитории с "
            f"контентом, что является критически важным индикатором эффективности коммуникационной "
            f"стратегии. Высокий уровень вовлечений свидетельствует о релевантности контента "
            f"интересам и потребностям целевой аудитории, а также о способности авторов "
            f"генерировать контент, стимулирующий активное взаимодействие пользователей."
        )
        summary_run2 = summary_para2.add_run(summary_text2)
        summary_run2.font.size = Pt(11)

        doc.add_paragraph()

        summary_para3 = doc.add_paragraph()
        summary_text3 = (
            f"Средний показатель Presence Score (PS) по всем авторам составил {avg_ps:.1f} баллов, "
            f"что является интегральной оценкой силы присутствия в социальных медиа с учетом "
            f"перцентильного ранжирования ключевых метрик. Средний уровень вовлеченности по "
            f"просмотрам (ER_view) зафиксирован на уровне {format_percent(avg_er * 100)}, что "
            f"представляет собой важнейший показатель качества контента и его способности "
            f"генерировать активное взаимодействие с аудиторией. Эти агрегированные показатели "
            f"позволяют оценить общую эффективность медиа-присутствия и сравнить текущее состояние "
            f"с отраслевыми бенчмарками и показателями предыдущих периодов."
        )
        summary_run3 = summary_para3.add_run(summary_text3)
        summary_run3.font.size = Pt(11)

        doc.add_paragraph()

        # Таблица ключевых метрик
        table = doc.add_table(rows=6, cols=2)
        table.style = "Light Grid Accent 1"

        metrics_data = [
            ("Общее количество авторов", format_number(total_authors)),
            ("Общая аудитория", format_number(total_followers)),
            ("Публикаций всего", format_number(total_posts)),
            ("Просмотров всего", format_number(total_views)),
            ("Средний Presence Score", f"{avg_ps:.1f} баллов"),
            ("Средний ER (по просмотрам)", format_percent(avg_er * 100)),
        ]

        for idx, (metric, value) in enumerate(metrics_data):
            row = table.rows[idx]
            row.cells[0].text = metric
            row.cells[1].text = value

            # Форматирование
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

            # Жирный шрифт для значений
            for paragraph in row.cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        doc.add_paragraph()

        # Динамика изменений
        if "previous_period" in data:
            doc.add_heading(
                "Сравнительный анализ динамики: текущий период vs предыдущий период",
                level=2,
            )

            # Собираем дельты
            total_delta_followers = 0
            total_delta_posts = 0
            total_delta_views = 0
            total_delta_engagement = 0

            prev_total_followers = 0
            prev_total_posts = 0
            prev_total_views = 0
            prev_total_engagement = 0

            for platform_data in platforms.values():
                deltas = platform_data["aggregated"].get("deltas", {})

                total_delta_followers += deltas.get("followers", {}).get("absolute", 0)
                total_delta_posts += deltas.get("posts", {}).get("absolute", 0)
                total_delta_views += deltas.get("views", {}).get("absolute", 0)
                total_delta_engagement += deltas.get("engagement", {}).get(
                    "absolute", 0
                )

            # Вычисляем предыдущие значения
            prev_total_followers = total_followers - total_delta_followers
            prev_total_posts = total_posts - total_delta_posts
            prev_total_views = total_views - total_delta_views
            prev_total_engagement = total_engagement - total_delta_engagement

            # Процентные изменения
            delta_followers_pct = (
                (total_delta_followers / prev_total_followers * 100)
                if prev_total_followers > 0
                else 0
            )
            delta_posts_pct = (
                (total_delta_posts / prev_total_posts * 100)
                if prev_total_posts > 0
                else 0
            )
            delta_views_pct = (
                (total_delta_views / prev_total_views * 100)
                if prev_total_views > 0
                else 0
            )
            delta_engagement_pct = (
                (total_delta_engagement / prev_total_engagement * 100)
                if prev_total_engagement > 0
                else 0
            )

            dynamics_para1 = doc.add_paragraph()
            dynamics_text1 = (
                f"Сравнительный анализ показателей текущего отчетного периода с аналогичными "
                f"показателями предыдущего периода позволяет выявить значимые тренды и тенденции "
                f"в развитии медиа-присутствия. В предыдущем периоде общая аудитория составляла "
                f"{format_number(prev_total_followers)} подписчиков, в то время как в текущем "
                f"периоде этот показатель достиг {format_number(total_followers)} подписчиков. "
                f"Таким образом, аудитория показала {get_trend_description(delta_followers_pct)} "
                f"с изменением {format_delta(delta_followers_pct, total_delta_followers)}, что "
                f"{'свидетельствует о положительной динамике привлечения новых подписчиков и росте медийного влияния' if delta_followers_pct > 0 else 'требует анализа причин оттока аудитории и корректировки контент-стратегии' if delta_followers_pct < 0 else 'указывает на стабильность аудитории при отсутствии значимых изменений'}."
            )
            dynamics_run1 = dynamics_para1.add_run(dynamics_text1)
            dynamics_run1.font.size = Pt(11)

            doc.add_paragraph()

            dynamics_para2 = doc.add_paragraph()
            dynamics_text2 = (
                f"Что касается публикационной активности, в предыдущем периоде было опубликовано "
                f"{format_number(prev_total_posts)} {'публикация' if prev_total_posts == 1 else 'публикаций'}, "
                f"тогда как в текущем периоде количество публикаций составило "
                f"{format_number(total_posts)} постов. Количество публикаций продемонстрировало "
                f"{get_trend_description(delta_posts_pct)} ({format_delta(delta_posts_pct, total_delta_posts)}). "
                f"{'Увеличение публикационной активности свидетельствует о более интенсивной контент-стратегии и может способствовать росту охвата и вовлеченности аудитории' if delta_posts_pct > 0 else 'Снижение публикационной активности может быть обусловлено стратегическим фокусом на качестве контента, а не количестве, либо внешними факторами, ограничивающими производство контента' if delta_posts_pct < 0 else 'Стабильность публикационной активности указывает на устойчивую контент-стратегию с регулярным графиком выпуска материалов'}. "
                f"Данный показатель необходимо рассматривать в комплексе с метриками эффективности "
                f"отдельных публикаций для получения полной картины результативности контента."
            )
            dynamics_run2 = dynamics_para2.add_run(dynamics_text2)
            dynamics_run2.font.size = Pt(11)

            doc.add_paragraph()

            dynamics_para3 = doc.add_paragraph()
            dynamics_text3 = (
                f"Анализ охвата контента демонстрирует следующую картину: в предыдущем периоде "
                f"совокупное количество просмотров составило {format_number(prev_total_views)}, "
                f"в то время как в текущем периоде этот показатель достиг "
                f"{format_number(total_views)} просмотров. Таким образом, показатель просмотров "
                f"показал {get_trend_description(delta_views_pct)} "
                f"({format_delta(delta_views_pct, total_delta_views)}). "
                f"{'Рост просмотров является критически важным позитивным сигналом, указывающим на расширение охвата аудитории, повышение видимости контента в алгоритмах социальных платформ, а также растущий интерес пользователей к публикуемым материалам' if delta_views_pct > 0 else 'Снижение просмотров требует детального анализа причин: это может быть связано с изменениями алгоритмов платформ, снижением качества или релевантности контента, либо возросшей конкуренцией в нише' if delta_views_pct < 0 else 'Стабильность просмотров при неизменной или растущей публикационной активности может свидетельствовать о достижении определенного потолка органического охвата'}. "
                f"Этот показатель является одним из ключевых индикаторов эффективности медиа-стратегии "
                f"и требует постоянного мониторинга."
            )
            dynamics_run3 = dynamics_para3.add_run(dynamics_text3)
            dynamics_run3.font.size = Pt(11)

            doc.add_paragraph()

            dynamics_para4 = doc.add_paragraph()
            dynamics_text4 = (
                f"Показатель вовлеченности аудитории претерпел значительные изменения между "
                f"периодами. В предыдущем периоде общее количество вовлечений (включая лайки, "
                f"комментарии, репосты и сохранения) составило {format_number(prev_total_engagement)}, "
                f"в текущем же периоде зафиксировано {format_number(total_engagement)} вовлечений. "
                f"Вовлеченность продемонстрировала {get_trend_description(delta_engagement_pct)} "
                f"({format_delta(delta_engagement_pct, total_delta_engagement)}). "
                f"{'Положительная динамика вовлеченности является наиболее ценным показателем, так как свидетельствует не просто о пассивном потреблении контента, а об активном взаимодействии аудитории с материалами, что критически важно для алгоритмов социальных платформ и органического роста охвата' if delta_engagement_pct > 0 else 'Снижение вовлеченности при сохранении или росте просмотров может указывать на ослабление эмоционального отклика аудитории, необходимость обновления контент-стратегии или диверсификации форматов публикаций' if delta_engagement_pct < 0 else 'Стабильность показателя вовлеченности при изменяющихся объемах охвата требует детального анализа качественных характеристик аудитории и типов контента'}. "
                f"Вовлеченность остается важнейшим показателем качества взаимодействия с аудиторией "
                f"и должна находиться в фокусе внимания при оптимизации контент-стратегии."
            )
            dynamics_run4 = dynamics_para4.add_run(dynamics_text4)
            dynamics_run4.font.size = Pt(11)

        doc.add_paragraph()

        # Распределение по платформам
        doc.add_heading("Распределение по платформам", level=2)

        for platform_key, platform_data in platforms.items():
            platform_name = get_platform_name(platform_key)
            agg = platform_data["aggregated"]

            platform_para = doc.add_paragraph()
            platform_run = platform_para.add_run(f"• {platform_name}: ")
            platform_run.font.bold = True
            platform_run.font.size = Pt(11)

            stats_run = platform_para.add_run(
                f"{format_number(agg['total_followers'])} подписчиков, "
                f"{format_number(agg['total_posts'])} постов, "
                f"{format_number(agg['total_views'])} просмотров, "
                f"Presence Score: {agg['avg_PS']:.1f}"
            )
            stats_run.font.size = Pt(11)

        # Разрыв страницы
        doc.add_page_break()
