"""
Блок титульной страницы отчета
"""

from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict


class CoverPageBlock:
    """Генерирует титульную страницу отчета"""

    @staticmethod
    def add_to_document(doc: Document, data: Dict) -> None:
        """
        Добавляет титульную страницу в документ

        Args:
            doc: Документ Word
            data: Данные отчета
        """
        # Заголовок отчета
        title = doc.add_heading("АНАЛИТИЧЕСКИЙ ОТЧЕТ", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(0, 102, 204)

        # Подзаголовок
        subtitle = doc.add_heading("Сравнительный анализ эффективности", level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.color.rgb = RGBColor(64, 64, 64)

        # Пустое пространство
        doc.add_paragraph()
        doc.add_paragraph()

        # Информация о периоде
        period_info = doc.add_paragraph()
        period_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

        period_data = data.get("period", {})
        start_date = datetime.fromisoformat(period_data["start"]).strftime("%d.%m.%Y")
        end_date = datetime.fromisoformat(period_data["end"]).strftime("%d.%m.%Y")
        days = period_data.get("days", 0)

        period_run = period_info.add_run(f"Период анализа: {start_date} — {end_date}\n")
        period_run.font.size = Pt(14)
        period_run.font.bold = True

        days_run = period_info.add_run(f"Длительность: {days} дней")
        days_run.font.size = Pt(12)

        # Информация о предыдущем периоде
        if "previous_period" in data:
            doc.add_paragraph()
            prev_info = doc.add_paragraph()
            prev_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

            prev_data = data["previous_period"]
            prev_start = datetime.fromisoformat(prev_data["start"]).strftime("%d.%m.%Y")
            prev_end = datetime.fromisoformat(prev_data["end"]).strftime("%d.%m.%Y")

            prev_run = prev_info.add_run(f"Период сравнения: {prev_start} — {prev_end}")
            prev_run.font.size = Pt(11)
            prev_run.font.italic = True
            prev_run.font.color.rgb = RGBColor(128, 128, 128)

        # Пустое пространство
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()

        # Платформы
        platforms_para = doc.add_paragraph()
        platforms_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        platforms_title = platforms_para.add_run("Анализируемые платформы:\n")
        platforms_title.font.size = Pt(12)
        platforms_title.font.bold = True

        platforms = data.get("platforms", {})
        from ..utils import get_platform_name

        for platform_key in platforms.keys():
            platform_name = get_platform_name(platform_key)
            platform_run = platforms_para.add_run(f"• {platform_name}\n")
            platform_run.font.size = Pt(11)

        # Пустое пространство
        doc.add_paragraph()
        doc.add_paragraph()

        # Дата формирования отчета
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        footer_run = footer_para.add_run(
            f"\nОтчет сформирован: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
        )
        footer_run.font.size = Pt(10)
        footer_run.font.italic = True
        footer_run.font.color.rgb = RGBColor(128, 128, 128)

        # Разрыв страницы
        doc.add_page_break()
