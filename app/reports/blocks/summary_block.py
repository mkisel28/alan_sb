"""
Блок общей сводки по всем платформам
"""

from docx import Document
from docx.shared import Pt, RGBColor
from typing import Dict
from ..utils import (
    format_number,
    format_percent,
    format_delta,
    get_platform_name,
    get_trend_description,
)


class SummaryBlock:
    """Генерирует общую сводку по всем платформам"""

    @staticmethod
    def add_to_document(doc: Document, data: Dict) -> None:
        """
        Добавляет общую сводку в документ

        Args:
            doc: Документ Word
            data: Данные отчета
        """
        # Заголовок сводки
        heading = doc.add_heading("ИСПОЛНИТЕЛЬНОЕ РЕЗЮМЕ", level=1)
        heading_run = heading.runs[0]
        heading_run.font.color.rgb = RGBColor(0, 102, 204)

        doc.add_paragraph()

        # Вводный текст
        intro = doc.add_paragraph()
        intro_run = intro.add_run(
            "Данный отчет представляет комплексный анализ эффективности присутствия "
            "в социальных сетях за отчетный период. Анализ охватывает ключевые показатели "
            "производительности (KPI), динамику развития аудитории и уровень вовлеченности "
            "пользователей на различных платформах."
        )
        intro_run.font.size = Pt(11)

        doc.add_paragraph()

        # Агрегированная статистика
        platforms = data.get("platforms", {})

        total_authors = sum(
            p["aggregated"]["total_authors"] for p in platforms.values()
        )
        total_followers = sum(
            p["aggregated"]["total_followers"] for p in platforms.values()
        )
        total_posts = sum(p["aggregated"]["total_posts"] for p in platforms.values())
        total_views = sum(p["aggregated"]["total_views"] for p in platforms.values())
        total_engagement = sum(
            p["aggregated"]["total_engagement"] for p in platforms.values()
        )

        # Средние значения
        avg_ps = (
            sum(p["aggregated"]["avg_PS"] for p in platforms.values()) / len(platforms)
            if platforms
            else 0
        )
        avg_er = (
            sum(p["aggregated"]["avg_ER_view"] for p in platforms.values())
            / len(platforms)
            if platforms
            else 0
        )

        # Ключевые показатели
        doc.add_heading("Ключевые показатели", level=2)

        summary_para = doc.add_paragraph()
        summary_text = (
            f"В течение отчетного периода была проанализирована активность {total_authors} "
            f"{'автора' if total_authors == 1 else 'авторов'} на {len(platforms)} "
            f"{'платформе' if len(platforms) == 1 else 'платформах'}. "
            f"Общая аудитория составила {format_number(total_followers)} подписчиков, "
            f"было опубликовано {format_number(total_posts)} "
            f"{'пост' if total_posts == 1 else 'постов'}, "
            f"которые собрали {format_number(total_views)} просмотров и "
            f"{format_number(total_engagement)} вовлечений (лайки, комментарии, репосты, сохранения)."
        )
        summary_run = summary_para.add_run(summary_text)
        summary_run.font.size = Pt(11)

        doc.add_paragraph()

        # Таблица ключевых метрик
        table = doc.add_table(rows=6, cols=2)
        table.style = "Light Grid Accent 1"

        metrics_data = [
            ("Общее количество авторов", format_number(total_authors)),
            ("Общая аудитория", format_number(total_followers)),
            ("Публикаций всего", format_number(total_posts)),
            ("Просмотров всего", format_number(total_views)),
            ("Средний Presence Score", f"{avg_ps:.1f} баллов"),
            ("Средний ER (по просмотрам)", format_percent(avg_er * 100)),
        ]

        for idx, (metric, value) in enumerate(metrics_data):
            row = table.rows[idx]
            row.cells[0].text = metric
            row.cells[1].text = value

            # Форматирование
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

            # Жирный шрифт для значений
            for paragraph in row.cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        doc.add_paragraph()

        # Динамика изменений
        if "previous_period" in data:
            doc.add_heading("Динамика изменений", level=2)

            # Собираем дельты
            total_delta_followers = 0
            total_delta_posts = 0
            total_delta_views = 0
            total_delta_engagement = 0

            prev_total_followers = 0
            prev_total_posts = 0
            prev_total_views = 0
            prev_total_engagement = 0

            for platform_data in platforms.values():
                deltas = platform_data["aggregated"].get("deltas", {})

                total_delta_followers += deltas.get("followers", {}).get("absolute", 0)
                total_delta_posts += deltas.get("posts", {}).get("absolute", 0)
                total_delta_views += deltas.get("views", {}).get("absolute", 0)
                total_delta_engagement += deltas.get("engagement", {}).get(
                    "absolute", 0
                )

            # Вычисляем предыдущие значения
            prev_total_followers = total_followers - total_delta_followers
            prev_total_posts = total_posts - total_delta_posts
            prev_total_views = total_views - total_delta_views
            prev_total_engagement = total_engagement - total_delta_engagement

            # Процентные изменения
            delta_followers_pct = (
                (total_delta_followers / prev_total_followers * 100)
                if prev_total_followers > 0
                else 0
            )
            delta_posts_pct = (
                (total_delta_posts / prev_total_posts * 100)
                if prev_total_posts > 0
                else 0
            )
            delta_views_pct = (
                (total_delta_views / prev_total_views * 100)
                if prev_total_views > 0
                else 0
            )
            delta_engagement_pct = (
                (total_delta_engagement / prev_total_engagement * 100)
                if prev_total_engagement > 0
                else 0
            )

            dynamics_para = doc.add_paragraph()
            dynamics_text = (
                f"По сравнению с предыдущим периодом наблюдается следующая динамика: "
                f"аудитория показала {get_trend_description(delta_followers_pct)} "
                f"({format_delta(delta_followers_pct, total_delta_followers)}), "
                f"количество публикаций - {get_trend_description(delta_posts_pct)} "
                f"({format_delta(delta_posts_pct, total_delta_posts)}), "
                f"просмотры - {get_trend_description(delta_views_pct)} "
                f"({format_delta(delta_views_pct, total_delta_views)}), "
                f"вовлеченность - {get_trend_description(delta_engagement_pct)} "
                f"({format_delta(delta_engagement_pct, total_delta_engagement)})."
            )
            dynamics_run = dynamics_para.add_run(dynamics_text)
            dynamics_run.font.size = Pt(11)

        doc.add_paragraph()

        # Распределение по платформам
        doc.add_heading("Распределение по платформам", level=2)

        for platform_key, platform_data in platforms.items():
            platform_name = get_platform_name(platform_key)
            agg = platform_data["aggregated"]

            platform_para = doc.add_paragraph()
            platform_run = platform_para.add_run(f"• {platform_name}: ")
            platform_run.font.bold = True
            platform_run.font.size = Pt(11)

            stats_run = platform_para.add_run(
                f"{format_number(agg['total_followers'])} подписчиков, "
                f"{format_number(agg['total_posts'])} постов, "
                f"{format_number(agg['total_views'])} просмотров, "
                f"Presence Score: {agg['avg_PS']:.1f}"
            )
            stats_run.font.size = Pt(11)

        # Разрыв страницы
        doc.add_page_break()
