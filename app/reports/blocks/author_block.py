"""
Блок данных по автору
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict
from ..utils import (
    format_number,
    format_percent,
    format_delta,
    interpret_ps,
    interpret_ms,
    interpret_er,
    get_trend_description,
)
from .charts_block import ChartsBlock


class AuthorBlock:
    """Генерирует блок с подробными данными по автору"""

    @staticmethod
    def add_to_document(doc: Document, author_data: Dict, platform_name: str) -> None:
        """
        Добавляет блок автора в документ

        Args:
            doc: Документ Word
            author_data: Данные автора
            platform_name: Название платформы
        """
        metrics = author_data["metrics"]
        scores = author_data["scores"]
        author_name = author_data["author_name"]
        username = author_data["username"]

        has_previous = "prev_metrics" in metrics
        prev_metrics = metrics.get("prev_metrics", {})

        # Заголовок автора
        heading = doc.add_heading(f"{author_name} (@{username})", level=3)
        heading_run = heading.runs[0]
        heading_run.font.color.rgb = RGBColor(0, 102, 204)

        # Основные метрики
        summary_para = doc.add_paragraph()
        summary_text = (
            f"{author_name} на платформе {platform_name} имеет аудиторию "
            f"в {format_number(metrics['F'])} подписчиков. "
            f"За отчетный период автором было опубликовано {metrics['P']} "
            f"{'пост' if metrics['P'] == 1 else 'постов' if metrics['P'] < 5 else 'постов'}, "
            f"которые собрали в совокупности {format_number(metrics['V'])} просмотров, "
            f"что составляет в среднем {format_number(int(metrics['V_avg']))} просмотров на одну публикацию. "
        )

        # Добавляем информацию о вовлеченности
        summary_text += (
            f"Общее количество вовлечений (лайки, комментарии, репосты и сохранения) "
            f"составило {format_number(metrics['E'])}, при этом уровень вовлеченности "
            f"по просмотрам (ER_view) достиг {format_percent(metrics['ER_view'] * 100)}, "
            f"что характеризуется как {interpret_er(metrics['ER_view'])}."
        )

        summary_run = summary_para.add_run(summary_text)
        summary_run.font.size = Pt(11)

        doc.add_paragraph()

        # Детальная таблица метрик
        doc.add_heading("Детальные метрики", level=4)

        table = doc.add_table(rows=9, cols=3 if has_previous else 2)
        table.style = "Light Grid Accent 1"

        # Заголовки
        headers = ["Метрика", "Текущий период"]
        if has_previous:
            headers.append("Предыдущий период")

        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)

        # Данные
        rows_data = [
            (
                "Подписчики (F)",
                format_number(metrics["F"]),
                format_number(prev_metrics.get("F", 0)) if has_previous else "",
            ),
            (
                "Публикации (P)",
                str(metrics["P"]),
                str(prev_metrics.get("P", 0)) if has_previous else "",
            ),
            (
                "Просмотры всего (V)",
                format_number(metrics["V"]),
                format_number(prev_metrics.get("V", 0)) if has_previous else "",
            ),
            (
                "Средние просмотры (V_avg)",
                format_number(int(metrics["V_avg"])),
                format_number(int(prev_metrics.get("V_avg", 0)))
                if has_previous
                else "",
            ),
            (
                "Вовлечения всего (E)",
                format_number(metrics["E"]),
                format_number(prev_metrics.get("E", 0)) if has_previous else "",
            ),
            (
                "ER по просмотрам",
                format_percent(metrics["ER_view"] * 100),
                format_percent(prev_metrics.get("ER_view", 0) * 100)
                if has_previous
                else "",
            ),
            (
                "Share Rate (SR)",
                format_percent(metrics["SR"] * 100, 3),
                format_percent(prev_metrics.get("SR", 0) * 100, 3)
                if has_previous
                else "",
            ),
            (
                "Comment Rate (CR)",
                format_percent(metrics["CR"] * 100, 3),
                format_percent(prev_metrics.get("CR", 0) * 100, 3)
                if has_previous
                else "",
            ),
        ]

        for idx, row_data in enumerate(rows_data, start=1):
            for col_idx, value in enumerate(row_data):
                table.rows[idx].cells[col_idx].text = value
                for paragraph in table.rows[idx].cells[col_idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        doc.add_paragraph()

        # Динамика
        if has_previous:
            doc.add_heading("Динамика изменений", level=4)

            dynamics_para = doc.add_paragraph()

            # Расчет изменений
            delta_f_pct = metrics["delta_F_percent"] * 100
            delta_v_pct = metrics.get("delta_V_avg_percent", 0) * 100
            delta_er_pct = metrics.get("delta_ER_percent", 0) * 100

            dynamics_text = (
                f"По сравнению с предыдущим периодом наблюдается "
                f"{get_trend_description(delta_f_pct)} аудитории "
                f"({format_delta(delta_f_pct, metrics['delta_F'])}). "
                f"Средние просмотры на пост показали {get_trend_description(delta_v_pct)} "
                f"({format_delta(delta_v_pct, is_percent=True)}). "
                f"Уровень вовлеченности изменился на {format_delta(delta_er_pct, is_percent=True)}, "
                f"что говорит о {get_trend_description(delta_er_pct).lower()}."
            )

            dynamics_run = dynamics_para.add_run(dynamics_text)
            dynamics_run.font.size = Pt(11)

            doc.add_paragraph()

        # Оценки
        doc.add_heading("Интегральные оценки", level=4)

        scores_para = doc.add_paragraph()

        ps_text = (
            f"Presence Score (PS) автора составляет {scores['PS']} баллов, "
            f'что соответствует категории "{interpret_ps(scores["PS"])}". '
            f"Данный показатель рассчитывается на основе перцентильного ранжирования "
            f"ключевых метрик присутствия относительно других авторов на платформе."
        )

        if has_previous and "prev_PS" in scores:
            ps_delta = scores["PS"] - scores["prev_PS"]
            ps_text += (
                f" В предыдущем периоде PS составлял {scores['prev_PS']} баллов, "
                f"таким образом изменение составило {format_delta(ps_delta, is_percent=False)} баллов."
            )

        scores_run = scores_para.add_run(ps_text)
        scores_run.font.size = Pt(11)

        doc.add_paragraph()

        # Momentum Score
        if "MS" in scores:
            ms_para = doc.add_paragraph()
            ms_text = (
                f"Momentum Score (MS) показывает {scores['MS']} баллов, "
                f"что характеризует {interpret_ms(scores['MS']).lower()} автора. "
                f"Этот индикатор отражает динамику развития канала, учитывая изменения "
                f"просмотров, вовлеченности и аудитории относительно конкурентов."
            )
            ms_run = ms_para.add_run(ms_text)
            ms_run.font.size = Pt(11)

            doc.add_paragraph()

        # График метрик
        try:
            chart_stream = ChartsBlock.create_metrics_comparison(
                author_name, metrics, prev_metrics if has_previous else None
            )
            doc.add_picture(chart_stream, width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            error_para = doc.add_paragraph()
            error_run = error_para.add_run(f"[График недоступен: {str(e)}]")
            error_run.font.italic = True
            error_run.font.color.rgb = RGBColor(255, 0, 0)

        doc.add_paragraph()
