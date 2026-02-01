"""
Блок данных по платформе
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict
from ..utils import (
    format_number,
    format_percent,
    format_delta,
    get_platform_name,
    get_trend_description,
)
from .author_block import AuthorBlock
from .charts_block import ChartsBlock


class PlatformBlock:
    """Генерирует блок данных по платформе"""

    @staticmethod
    def add_to_document(doc: Document, platform_key: str, platform_data: Dict) -> None:
        """
        Добавляет блок платформы в документ

        Args:
            doc: Документ Word
            platform_key: Ключ платформы
            platform_data: Данные платформы
        """
        platform_name = get_platform_name(platform_key)
        agg = platform_data["aggregated"]
        authors = platform_data["authors"]

        # Заголовок платформы
        heading = doc.add_heading(platform_name.upper(), level=1)
        heading_run = heading.runs[0]
        heading_run.font.color.rgb = RGBColor(0, 102, 204)

        doc.add_paragraph()

        # Обзор платформы
        overview_para = doc.add_paragraph()
        overview_text = (
            f"Анализ эффективности присутствия на платформе {platform_name} охватывает "
            f"{agg['total_authors']} "
            f"{'автора' if agg['total_authors'] == 1 else 'авторов'} "
            f"с совокупной аудиторией в {format_number(agg['total_followers'])} подписчиков. "
            f"За отчетный период на данной платформе было размещено "
            f"{format_number(agg['total_posts'])} публикаций, "
            f"которые в общей сложности получили {format_number(agg['total_views'])} просмотров "
            f"и {format_number(agg['total_engagement'])} вовлечений."
        )
        overview_run = overview_para.add_run(overview_text)
        overview_run.font.size = Pt(11)

        doc.add_paragraph()

        # Агрегированные метрики
        doc.add_heading("Агрегированные показатели", level=2)

        metrics_para = doc.add_paragraph()
        metrics_text = (
            f"Средний показатель Presence Score по платформе составляет {agg['avg_PS']:.1f} баллов, "
            f"а средний уровень вовлеченности (ER_view) находится на уровне "
            f"{format_percent(agg['avg_ER_view'] * 100)}. "
        )

        if "avg_MS" in agg:
            metrics_text += (
                f"Momentum Score в среднем составляет {agg['avg_MS']:.1f} баллов, "
                f"что характеризует общую динамику развития присутствия на платформе."
            )

        metrics_run = metrics_para.add_run(metrics_text)
        metrics_run.font.size = Pt(11)

        doc.add_paragraph()

        # Динамика
        if "deltas" in agg:
            doc.add_heading("Динамика по сравнению с предыдущим периодом", level=2)

            deltas = agg["deltas"]

            dynamics_para = doc.add_paragraph()
            dynamics_text = (
                f"Изменение ключевых показателей платформы:\n\n"
                f"• Подписчики: {get_trend_description(deltas['followers']['percent'])} - "
                f"{format_delta(deltas['followers']['percent'], deltas['followers']['absolute'])}\n"
                f"• Публикации: {get_trend_description(deltas['posts']['percent'])} - "
                f"{format_delta(deltas['posts']['percent'], deltas['posts']['absolute'])}\n"
                f"• Просмотры: {get_trend_description(deltas['views']['percent'])} - "
                f"{format_delta(deltas['views']['percent'], deltas['views']['absolute'])}\n"
                f"• Вовлеченность: {get_trend_description(deltas['engagement']['percent'])} - "
                f"{format_delta(deltas['engagement']['percent'], deltas['engagement']['absolute'])}\n"
                f"• Engagement Rate: {get_trend_description(deltas['ER_view']['percent'])} - "
                f"{format_delta(deltas['ER_view']['percent'], is_percent=True)}"
            )
            dynamics_run = dynamics_para.add_run(dynamics_text)
            dynamics_run.font.size = Pt(10)

            doc.add_paragraph()

        # Графики по платформе
        try:
            # График средних просмотров
            labels = [a["author_name"] for a in authors]
            v_avg_values = [a["metrics"]["V_avg"] for a in authors]

            chart_stream = ChartsBlock.create_bar_chart(
                labels,
                v_avg_values,
                f"Средние просмотры на пост - {platform_name}",
                "Просмотры",
                color="#4A90E2",
            )
            doc.add_picture(chart_stream, width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()

            # График ER
            er_values = [a["metrics"]["ER_view"] * 100 for a in authors]

            chart_stream = ChartsBlock.create_bar_chart(
                labels,
                er_values,
                f"Engagement Rate - {platform_name}",
                "ER (%)",
                color="#67C23A",
            )
            doc.add_picture(chart_stream, width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()

        except Exception as e:
            error_para = doc.add_paragraph()
            error_run = error_para.add_run(f"[Графики недоступны: {str(e)}]")
            error_run.font.italic = True
            error_run.font.color.rgb = RGBColor(255, 0, 0)

        # Детальные данные по каждому автору
        doc.add_heading("Детальный анализ по авторам", level=2)

        for idx, author_data in enumerate(authors, start=1):
            AuthorBlock.add_to_document(doc, author_data, platform_name)

            # Разделитель между авторами
            if idx < len(authors):
                doc.add_paragraph("_" * 70)
                doc.add_paragraph()

        # Разрыв страницы после платформы
        doc.add_page_break()
